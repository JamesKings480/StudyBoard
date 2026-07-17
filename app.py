from flask import Flask, render_template, redirect, url_for, flash, request, jsonify, send_file
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_wtf.csrf import CSRFProtect
from datetime import date, timedelta
from config import Config
import io
import os
import pdfplumber
from docx import Document as DocxDocument
from werkzeug.utils import secure_filename
from groq_client import generate_subtasks

HSC_SUBJECTS = {
    'English Standard': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-standard-stage-6-2017',
    'English Advanced': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-advanced-stage-6-2017',
    'English Studies': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-studies-stage-6-2017',
    'English EAL/D': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-eald-stage-6-2017',
    'English Extension 1': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-extension-stage-6-2017',
    'English Extension 2': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/english/english-extension-stage-6-2017',
    'Mathematics Standard': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/mathematics/mathematics-standard-stage-6-2017',
    'Mathematics Advanced': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/mathematics/mathematics-advanced-stage-6-2017',
    'Mathematics Extension 1': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/mathematics/mathematics-extension-1-stage-6-2017',
    'Mathematics Extension 2': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/mathematics/mathematics-extension-2-stage-6-2017',
    'Drama': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/creative-arts/drama-stage-6-2009',
    'Business Studies': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/business-studies-stage-6-2010',
    'Economics': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/economics-stage-6-2009',
    'Geography': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/geography-stage-6-2009',
    'Aboriginal Studies': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/aboriginal-studies-stage-6-2010',
    'Ancient History': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/ancient-history-stage-6-2017',
    'Modern History': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/modern-history-stage-6-2017',
    'History Extension': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/history-extension-stage-6-2017',
    'Legal Studies': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/legal-studies-stage-6-2009',
    'French Continuers': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/languages/french-stage-6-2009',
    'Chinese Continuers': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/languages/chinese-stage-6-2009',
    'Latin Continuers': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/languages/latin-stage-6-2010',
    'Modern Greek Beginners': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/languages/modern-greek-stage-6-2009',
    'Music 1': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/creative-arts/music-stage-6-2009',
    'Music 2': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/creative-arts/music-stage-6-2009',
    'Music Extension': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/creative-arts/music-extension-stage-6-2009',
    'Health and Movement Science': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/pdhpe/health-and-movement-science-stage-6-2024',
    'Studies of Religion I': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/studies-of-religion-stage-6-2009',
    'Studies of Religion II': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/hsie/studies-of-religion-stage-6-2009',
    'Biology': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/science/biology-stage-6-2017',
    'Chemistry': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/science/chemistry-stage-6-2017',
    'Earth and Environmental Science': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/science/earth-and-environmental-science-stage-6-2017',
    'Physics': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/science/physics-stage-6-2017',
    'Science Extension': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/science/science-extension-stage-6-2017',
    'Design and Technology': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/tas/design-and-technology-stage-6-2018',
    'Engineering Studies': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/tas/engineering-studies-stage-6-2022',
    'Industrial Technology': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/tas/industrial-technology-stage-6-2019',
    'Software Engineering': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/tas/software-engineering-11-12-2022',
    'Visual Arts': 'https://www.nsw.gov.au/education-and-training/nesa/curriculum/creative-arts/visual-arts-stage-6-2016',
}

from models import db, User, Subject, Assessment, Task, StudySession, TodoItem, SubjectFile, Topic, Flashcard, FlashcardReview
from forms import RegistrationForm, LoginForm, SubjectForm, AssessmentForm, MarkForm, TaskForm, TodoItemForm, FlashcardForm

app = Flask(__name__)
app.config.from_object(Config)
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_NOTIFICATION_CHARS = 20000
FILE_MIMETYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'doc': 'application/msword',
    'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'txt': 'text/plain',
    'png': 'image/png',
    'jpg': 'image/jpeg',
    'jpeg': 'image/jpeg',
}
OPEN_IN_BROWSER = {'pdf', 'txt', 'png', 'jpg', 'jpeg'}
NOTE_EXTENSIONS = {'pdf', 'docx', 'doc', 'pptx', 'txt', 'png', 'jpg', 'jpeg'}
MAX_NOTE_SIZE = 10 * 1024 * 1024
STRONG_ACCURACY_PERCENT = 70
MIN_REVIEWS_FOR_STRONG = 3
MAX_FILE_SIZE = 5 * 1024 * 1024 

def allowed_file(filename, allowed=None):
    allowed = allowed or ALLOWED_EXTENSIONS
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed


def safe_filename(original_name):
    filename = secure_filename(original_name)
    if '.' not in filename:
        filename = 'file.' + original_name.rsplit('.', 1)[1].lower()
    return filename


def serve_stored_file(name, data):
    ext = name.rsplit('.', 1)[-1].lower()
    return send_file(
        io.BytesIO(data),
        mimetype=FILE_MIMETYPES.get(ext, 'application/octet-stream'),
        as_attachment=ext not in OPEN_IN_BROWSER,
        download_name=name
    )

def extract_text_from_file(uploaded_file, filename):
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    text = ''
    try:
        uploaded_file.seek(0)
        file_bytes = io.BytesIO(uploaded_file.read())
        if ext == 'pdf':
            with pdfplumber.open(file_bytes) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
        elif ext == 'docx':
            document = DocxDocument(file_bytes)
            for para in document.paragraphs:
                text += para.text + '\n'
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + '\n'
        elif ext == 'txt':
            text = file_bytes.read().decode('utf-8', errors='ignore')
    except Exception as error:
        print('EXTRACTION ERROR [' + ext + ']:', error)
        return ''
    return text.strip()


def get_notification_text(form):
    typed_text = form.task_notification.data.strip() if form.task_notification.data else ''

    uploaded_file = request.files.get('task_file')
    if not uploaded_file or not uploaded_file.filename:
        return typed_text, None, None

    original_name = uploaded_file.filename
    if original_name.lower().endswith('.doc'):
        return typed_text, None, 'Old .doc files cannot be read. Open it in Word, save it as .docx, then upload again.'

    if not allowed_file(original_name):
        return typed_text, None, 'Only PDF, Word (.docx) and text files are allowed.'

    uploaded_file.seek(0, os.SEEK_END)
    file_size = uploaded_file.tell()
    uploaded_file.seek(0)
    if file_size > MAX_FILE_SIZE:
        return typed_text, None, 'File is too large. Maximum size is 5MB.'

    filename = safe_filename(original_name)

    file_data = uploaded_file.read()
    extracted = extract_text_from_file(uploaded_file, filename)
    if not extracted:
        return typed_text, None, 'Could not read any text from that file. It may be a scan rather than real text. Try another file or type the task in.'

    file_info = {'name': filename, 'size': file_size, 'data': file_data}
    return extracted[:MAX_NOTIFICATION_CHARS], file_info, None

FALLBACK_STEPS = {
    'Essay': [
        'Read the task notification and marking criteria',
        'Research and gather your sources',
        'Plan the structure and write your thesis',
        'Write the first draft',
        'Edit, proofread and check the referencing',
    ],
    'Exam': [
        'Collect all your notes and past papers',
        'Summarise each topic onto one page',
        'Practise questions under timed conditions',
        'Review everything you got wrong',
        'Final revision of your weakest topics',
    ],
    'Research Task': [
        'Read the task notification and marking criteria',
        'Find and record your sources',
        'Take notes and gather your evidence',
        'Write the draft',
        'Proofread and finish the bibliography',
    ],
    'Practical': [
        'Read the task notification and safety requirements',
        'Plan the method and gather materials',
        'Carry out the practical and record results',
        'Analyse your results',
        'Write up the report',
    ],
    'Presentation': [
        'Read the task notification and marking criteria',
        'Research the content',
        'Build the slides or visuals',
        'Rehearse out loud and time yourself',
        'Final rehearsal and fix the timing',
    ],
    'Assignment': [
        'Read the task notification and marking criteria',
        'Break the task into its required parts',
        'Do the research or the working',
        'Write or build the draft',
        'Check against the marking criteria and submit',
    ],
    'Other': [
        'Read the task notification and marking criteria',
        'Plan what needs doing',
        'Do the main body of the work',
        'Review against the marking criteria',
        'Final check and submit',
    ],
}


def build_fallback_subtasks(assessment_type, days_available):
    steps = FALLBACK_STEPS.get(assessment_type, FALLBACK_STEPS['Other'])
    subtasks = []
    for index, title in enumerate(steps):
        fraction = (len(steps) - index - 1) / len(steps)
        subtasks.append({'title': title, 'days_before_due': int(days_available * fraction)})
    return subtasks


def save_subtasks(assessment, subtasks):
    today = date.today()
    for item in subtasks:
        task_date = assessment.due_date - timedelta(days=item['days_before_due'])
        if task_date < today:
            task_date = today
        task = Task(
            title=item['title'],
            description=item.get('description') or None,
            scheduled_date=task_date,
            assessment_id=assessment.id
        )
        db.session.add(task)
    db.session.commit()

csrf = CSRFProtect(app)
db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    return response

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    form = RegistrationForm()
    if form.validate_on_submit():
        existing = User.query.filter_by(email=form.email.data.lower().strip()).first()
        if existing:
            flash('Unable to create account. Please try a different email.', 'danger')
            return redirect(url_for('register'))
        user = User(email=form.email.data.lower().strip())
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Account created successfully! Please log in.', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data.lower().strip()).first()
        if user is None or not user.check_password(form.password.data):
            flash('Invalid email or password.', 'danger')
            return redirect(url_for('login'))
        login_user(user)
        return redirect(url_for('dashboard'))
    return render_template('login.html', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    today = date.today()
    start_of_week = today - timedelta(days=today.weekday())
    end_of_week = start_of_week + timedelta(days=6)

    subjects = Subject.query.filter_by(user_id=current_user.id).all()
    upcoming_assessments = Assessment.query.join(Subject).filter(
        Subject.user_id == current_user.id,
        Assessment.due_date >= today,
        Assessment.status != 'Completed'
    ).order_by(Assessment.due_date).limit(10).all()
    todays_tasks = Task.query.join(Assessment).join(Subject).filter(
        Subject.user_id == current_user.id,
        Task.scheduled_date <= today,
        db.or_(Task.status == 'Incomplete', Task.scheduled_date == today)
    ).all()
    todays_todos = TodoItem.query.join(Subject).filter(
        Subject.user_id == current_user.id,
        TodoItem.scheduled_date <= today,
        db.or_(TodoItem.status == 'Incomplete', TodoItem.scheduled_date == today)
    ).all()
    todays_items = sorted(
        todays_tasks + todays_todos,
        key=lambda item: (item.status == 'Complete', item.scheduled_date)
    )

    sessions = StudySession.query.filter(
        StudySession.user_id == current_user.id,
        StudySession.date >= start_of_week,
        StudySession.date <= end_of_week
    ).all()
    total_weekly_minutes = sum(s.duration_minutes for s in sessions)
    total_seconds = int(total_weekly_minutes * 60)
    if total_seconds >= 3600:
        weekly_display = str(total_seconds // 3600) + 'h ' + str((total_seconds % 3600) // 60) + 'm'
    elif total_seconds >= 60:
        weekly_display = str(total_seconds // 60) + 'm ' + str(total_seconds % 60) + 's'
    else:
        weekly_display = str(total_seconds) + 's'

    daily_totals = [0] * 7
    for s in sessions:
        day_index = (s.date - start_of_week).days
        if 0 <= day_index < 7:
            daily_totals[day_index] += s.duration_minutes
    max_daily = max(daily_totals) if max(daily_totals) > 0 else 1
    grades = {s.id: get_grade_summary(s) for s in subjects}
    show_tour = not current_user.has_seen_tour
    todo_form = TodoItemForm()
    todo_form.subject_id.choices = [(s.id, s.name) for s in subjects]
    todo_form.scheduled_date.data = today
    day_names = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
    return render_template('dashboard.html', subjects=subjects, upcoming_assessments=upcoming_assessments, todays_items=todays_items, todo_form=todo_form, grades=grades, show_tour=show_tour, total_weekly_minutes=total_weekly_minutes, weekly_display=weekly_display, daily_totals=daily_totals, max_daily=max_daily, day_names=day_names, today=today)

@app.route('/tour/seen', methods=['POST'])
@login_required
def mark_tour_seen():
    current_user.has_seen_tour = True
    db.session.commit()
    return jsonify({'success': True})


@app.route('/study/save', methods=['POST'])
@login_required
def save_study_session():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'Invalid data'}), 400
    subject_id = data.get('subject_id')
    duration = data.get('duration_minutes', 0)

    subject = Subject.query.get(subject_id)
    if not subject or subject.user_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403

    session = StudySession(
        duration_minutes=round(duration, 2),
        date=date.today(),
        subject_id=subject_id,
        user_id=current_user.id
    )
    db.session.add(session)
    db.session.commit()
    return jsonify({'success': True, 'duration': session.duration_minutes})

@app.route('/subjects')
@login_required
def subjects_list():
    subjects = Subject.query.filter_by(user_id=current_user.id).order_by(Subject.name).all()
    return render_template('subjects.html', subjects=subjects)

@app.route('/subject/new', methods=['GET', 'POST'])
@login_required
def create_subject():
    form = SubjectForm()
    form.name.choices = [(s, s) for s in HSC_SUBJECTS.keys()]
    if form.validate_on_submit():
        subject = Subject(
            name=form.name.data,
            colour=form.colour.data or '#4A90D9',
            year_level=form.year_level.data,
            nesa_url=HSC_SUBJECTS.get(form.name.data),
            user_id=current_user.id
        )
        db.session.add(subject)
        db.session.commit()
        flash(f'Subject "{subject.name}" created successfully!', 'success')
        return redirect(url_for('subject_detail', subject_id=subject.id))
    return render_template('subject_form.html', form=form, title='Add New Subject')

@app.route('/subject/<int:subject_id>')
@login_required
def subject_detail(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    assessments = Assessment.query.filter_by(subject_id=subject.id).order_by(Assessment.due_date).all()
    return render_template('subject_detail.html', subject=subject, assessments=assessments,
                           grade=get_grade_summary(subject), today=date.today())

@app.route('/subject/<int:subject_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_subject(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = SubjectForm(obj=subject)
    form.name.choices = [(s, s) for s in HSC_SUBJECTS.keys()]
    if form.validate_on_submit():
        subject.name = form.name.data
        subject.colour = form.colour.data or '#4A90D9'
        subject.year_level = form.year_level.data
        subject.nesa_url = HSC_SUBJECTS.get(form.name.data)
        db.session.commit()
        flash(f'Subject "{subject.name}" updated!', 'success')
        return redirect(url_for('subject_detail', subject_id=subject.id))
    return render_template('subject_form.html', form=form, title='Edit Subject')


@app.route('/subject/<int:subject_id>/target', methods=['POST'])
@login_required
def set_target(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    try:
        target = float(request.form.get('target_grade', ''))
    except ValueError:
        flash('Enter a target between 1 and 100.', 'danger')
        return redirect(url_for('subject_detail', subject_id=subject.id))
    if not 1 <= target <= 100:
        flash('Enter a target between 1 and 100.', 'danger')
        return redirect(url_for('subject_detail', subject_id=subject.id))
    subject.target_grade = target
    db.session.commit()
    flash(f'Target for {subject.name} set to {target}%.', 'success')
    return redirect(url_for('subject_detail', subject_id=subject.id))


@app.route('/subject/<int:subject_id>/delete', methods=['POST'])
@login_required
def delete_subject(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    name = subject.name
    db.session.delete(subject)
    db.session.commit()
    flash(f'Subject "{name}" deleted.', 'info')
    return redirect(url_for('subjects_list'))

@app.route('/subject/<int:subject_id>/file/new', methods=['POST'])
@login_required
def upload_subject_file(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    uploaded_file = request.files.get('subject_file')
    if not uploaded_file or not uploaded_file.filename:
        flash('Choose a file to upload first.', 'warning')
        return redirect(url_for('subject_detail', subject_id=subject.id))

    if not allowed_file(uploaded_file.filename, NOTE_EXTENSIONS):
        flash('You can upload PDF, Word, PowerPoint, text and image files.', 'danger')
        return redirect(url_for('subject_detail', subject_id=subject.id))

    uploaded_file.seek(0, os.SEEK_END)
    file_size = uploaded_file.tell()
    uploaded_file.seek(0)
    if file_size > MAX_NOTE_SIZE:
        flash('That file is too large. Maximum size is 10MB.', 'danger')
        return redirect(url_for('subject_detail', subject_id=subject.id))

    note = SubjectFile(
        file_name=safe_filename(uploaded_file.filename),
        file_size=file_size,
        file_data=uploaded_file.read(),
        subject_id=subject.id
    )
    db.session.add(note)
    db.session.commit()
    flash(f'"{note.file_name}" uploaded.', 'success')
    return redirect(url_for('subject_detail', subject_id=subject.id))


@app.route('/subject/file/<int:file_id>')
@login_required
def subject_file(file_id):
    note = SubjectFile.query.get_or_404(file_id)
    if note.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    return serve_stored_file(note.file_name, note.file_data)


@app.route('/subject/file/<int:file_id>/delete', methods=['POST'])
@login_required
def delete_subject_file(file_id):
    note = SubjectFile.query.get_or_404(file_id)
    if note.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    subject_id = note.subject_id
    db.session.delete(note)
    db.session.commit()
    flash('File deleted.', 'info')
    return redirect(url_for('subject_detail', subject_id=subject_id))


@app.route('/subject/<int:subject_id>/assessment/new', methods=['GET', 'POST'])
@login_required
def create_assessment(subject_id):
    subject = Subject.query.get_or_404(subject_id)
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = AssessmentForm()
    if form.validate_on_submit():
        notification_text, file_info, upload_error = get_notification_text(form)
        if upload_error:
            flash(upload_error, 'danger')
            return render_template('assessment_form.html', form=form, subject=subject, assessment=None, title='New Assessment')

        assessment = Assessment(
            name=form.name.data.strip(),
            due_date=form.due_date.data,
            weighting=form.weighting.data,
            assessment_type=form.assessment_type.data,
            task_notification=notification_text,
            subject_id=subject.id
        )
        if file_info:
            assessment.task_file_name = file_info['name']
            assessment.task_file_size = file_info['size']
            assessment.task_file_data = file_info['data']
        db.session.add(assessment)
        db.session.commit()

        days_available = (assessment.due_date - date.today()).days
        subtasks = generate_subtasks(assessment.assessment_type, days_available, assessment.task_notification)
        if subtasks:
            save_subtasks(assessment, subtasks)
            flash(f'Assessment "{assessment.name}" created with an AI study plan!', 'success')
        else:
            save_subtasks(assessment, build_fallback_subtasks(assessment.assessment_type, days_available))
            flash(f'Assessment "{assessment.name}" created, but the AI plan was not available so I have used a standard {assessment.assessment_type.lower()} plan. You can edit the subtasks below.', 'warning')
        return redirect(url_for('assessment_detail', assessment_id=assessment.id))
    return render_template('assessment_form.html', form=form, subject=subject, assessment=None, title='New Assessment')

@app.route('/assessment/<int:assessment_id>')
@login_required
def assessment_detail(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    subject = assessment.subject
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    tasks = Task.query.filter_by(assessment_id=assessment.id).order_by(Task.scheduled_date).all()
    mark_form = MarkForm()
    return render_template('assessment_detail.html', assessment=assessment, subject=subject, tasks=tasks, mark_form=mark_form, today=date.today())

@app.route('/assessment/<int:assessment_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_assessment(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    subject = assessment.subject
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = AssessmentForm(obj=assessment)
    if request.method == 'GET' and assessment.task_file_name:
        form.task_notification.data = ''

    if form.validate_on_submit():
        notification_text, file_info, upload_error = get_notification_text(form)
        if upload_error:
            flash(upload_error, 'danger')
            return render_template('assessment_form.html', form=form, subject=subject, assessment=assessment, title='Edit Assessment')

        assessment.name = form.name.data.strip()
        assessment.due_date = form.due_date.data
        assessment.weighting = form.weighting.data
        assessment.assessment_type = form.assessment_type.data

        if file_info:
            assessment.task_notification = notification_text
            assessment.task_file_name = file_info['name']
            assessment.task_file_size = file_info['size']
            assessment.task_file_data = file_info['data']
        elif request.form.get('remove_task_file') == '1':
            assessment.task_notification = notification_text
            assessment.task_file_name = None
            assessment.task_file_size = None
            assessment.task_file_data = None
        elif not assessment.task_file_name:
            assessment.task_notification = notification_text

        db.session.commit()
        flash(f'Assessment "{assessment.name}" updated!', 'success')
        return redirect(url_for('assessment_detail', assessment_id=assessment.id))
    return render_template('assessment_form.html', form=form, subject=subject, assessment=assessment, title='Edit Assessment')


@app.route('/assessment/<int:assessment_id>/file')
@login_required
def assessment_file(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    if assessment.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    if not assessment.task_file_data:
        flash('There is no file attached to that assessment.', 'warning')
        return redirect(url_for('assessment_detail', assessment_id=assessment.id))

    return serve_stored_file(assessment.task_file_name, assessment.task_file_data)


@app.route('/assessment/<int:assessment_id>/delete', methods=['POST'])
@login_required
def delete_assessment(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    subject = assessment.subject
    if subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    subject_id = subject.id
    db.session.delete(assessment)
    db.session.commit()
    flash('Assessment deleted.', 'info')
    return redirect(url_for('subject_detail', subject_id=subject_id))

@app.route('/assessment/<int:assessment_id>/mark', methods=['POST'])
@login_required
def save_mark(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    if assessment.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = MarkForm()
    if form.validate_on_submit():
        assessment.mark = form.mark.data
        assessment.status = 'Completed'
        db.session.commit()
        flash(f'Mark saved: {assessment.mark}%', 'success')
    return redirect(url_for('assessment_detail', assessment_id=assessment.id))

@app.route('/task/<int:task_id>/toggle', methods=['POST'])
@login_required
def toggle_task(task_id):
    task = Task.query.get_or_404(task_id)
    if task.assessment.subject.user_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403
    task.status = 'Complete' if task.status == 'Incomplete' else 'Incomplete'
    db.session.commit()
    return redirect(request.referrer or url_for('dashboard'))

def _review_counts_by_topic(user_id, only_correct=False):
    query = db.session.query(
        Flashcard.topic_id, db.func.count(FlashcardReview.id)
    ).select_from(FlashcardReview).join(
        Flashcard, FlashcardReview.flashcard_id == Flashcard.id
    ).join(
        Topic, Flashcard.topic_id == Topic.id
    ).join(
        Subject, Topic.subject_id == Subject.id
    ).filter(Subject.user_id == user_id)
    if only_correct:
        query = query.filter(FlashcardReview.was_correct == True)
    return dict(query.group_by(Flashcard.topic_id).all())


def get_topic_stats(user_id, subject_id=None):
    topic_query = Topic.query.join(Subject).filter(Subject.user_id == user_id)
    if subject_id:
        topic_query = topic_query.filter(Topic.subject_id == subject_id)
    topics = topic_query.order_by(Subject.name, Topic.name).all()

    card_counts = dict(db.session.query(
        Flashcard.topic_id, db.func.count(Flashcard.id)
    ).join(Topic, Flashcard.topic_id == Topic.id
    ).join(Subject, Topic.subject_id == Subject.id
    ).filter(Subject.user_id == user_id
    ).group_by(Flashcard.topic_id).all())

    total_reviews = _review_counts_by_topic(user_id)
    correct_reviews = _review_counts_by_topic(user_id, only_correct=True)

    strong = []
    work_on = []
    for topic in topics:
        cards = card_counts.get(topic.id, 0)
        if cards == 0:
            continue
        reviews = total_reviews.get(topic.id, 0)
        correct = correct_reviews.get(topic.id, 0)
        accuracy = round(correct / reviews * 100) if reviews else None
        stat = {'topic': topic, 'cards': cards, 'reviews': reviews, 'accuracy': accuracy}
        if reviews >= MIN_REVIEWS_FOR_STRONG and accuracy >= STRONG_ACCURACY_PERCENT:
            strong.append(stat)
        else:
            work_on.append(stat)

    strong.sort(key=lambda s: -s['accuracy'])
    work_on.sort(key=lambda s: (s['accuracy'] is not None, s['accuracy'] or 0))
    return strong, work_on


def calculate_weighted_grade(subject_id):
    assessments = Assessment.query.filter_by(subject_id=subject_id, status='Completed').all()
    weighted_sum = 0
    total_weight = 0
    for a in assessments:
        if a.mark is None:
            continue
        weighted_sum += (a.mark / 100) * a.weighting
        total_weight += a.weighting
    if total_weight == 0:
        return 0
    return (weighted_sum / total_weight) * 100


def calculate_required_mark(subject_id, target_grade, current_grade):
    completed = Assessment.query.filter_by(subject_id=subject_id, status='Completed').all()
    completed_weight = sum(a.weighting for a in completed if a.mark is not None)
    remaining_weight = 100 - completed_weight
    if remaining_weight <= 0:
        return -1 if current_grade >= target_grade else 101
    earned_weighted_marks = (current_grade / 100) * completed_weight
    needed_weighted_marks = target_grade - earned_weighted_marks
    return round((needed_weighted_marks / remaining_weight) * 100, 1)


def get_grade_summary(subject):
    completed = Assessment.query.filter_by(subject_id=subject.id, status='Completed').all()
    completed_weight = sum(a.weighting for a in completed if a.mark is not None)
    entered_weight = sum(a.weighting for a in Assessment.query.filter_by(subject_id=subject.id).all())
    target = subject.target_grade or 80
    current = calculate_weighted_grade(subject.id)
    required = calculate_required_mark(subject.id, target, current)
    return {
        'has_marks': completed_weight > 0,
        'current': round(current, 1),
        'target': target,
        'required': required,
        'completed_weight': round(completed_weight, 1),
        'remaining_weight': round(max(0, 100 - completed_weight), 1),
        'entered_weight': round(entered_weight, 1),
        'achieved': required < 0,
        'impossible': required > 100,
    }


def get_or_create_topic(subject, raw_name):
    clean = ' '.join(raw_name.split())
    topic = Topic.query.filter(
        Topic.subject_id == subject.id,
        db.func.lower(Topic.name) == clean.lower()
    ).first()
    if topic:
        return topic
    topic = Topic(name=clean, subject_id=subject.id)
    db.session.add(topic)
    db.session.flush()
    return topic


def build_flashcard_form():
    subjects = Subject.query.filter_by(user_id=current_user.id).order_by(Subject.name).all()
    form = FlashcardForm()
    form.subject_id.choices = [(s.id, s.name) for s in subjects]
    return form, subjects


def flash_form_errors(form):
    for field_errors in form.errors.values():
        for error in field_errors:
            flash(error, 'danger')


@app.route('/flashcard/<int:card_id>/review', methods=['POST'])
@login_required
def record_review(card_id):
    card = Flashcard.query.get_or_404(card_id)
    if card.topic.subject.user_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403
    data = request.get_json(silent=True) or {}
    if 'was_correct' not in data:
        return jsonify({'error': 'No result sent'}), 400
    review = FlashcardReview(was_correct=bool(data['was_correct']), flashcard_id=card.id)
    db.session.add(review)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/flashcards')
@login_required
def flashcards_page():
    form, subjects = build_flashcard_form()
    topics = Topic.query.join(Subject).filter(
        Subject.user_id == current_user.id
    ).order_by(Subject.name, Topic.name).all()
    strong, work_on = get_topic_stats(current_user.id)
    return render_template('flashcards.html', form=form, subjects=subjects, topics=topics,
                           strong=strong, work_on=work_on,
                           strong_threshold=STRONG_ACCURACY_PERCENT,
                           min_reviews=MIN_REVIEWS_FOR_STRONG)


@app.route('/flashcard/new', methods=['POST'])
@login_required
def create_flashcard():
    form, subjects = build_flashcard_form()
    if not form.validate_on_submit():
        flash_form_errors(form)
        return redirect(url_for('flashcards_page'))

    subject = Subject.query.get(form.subject_id.data)
    if not subject or subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    topic = get_or_create_topic(subject, form.topic_name.data)
    card = Flashcard(
        question=form.question.data.strip(),
        answer=form.answer.data.strip(),
        topic_id=topic.id
    )
    db.session.add(card)
    db.session.commit()
    flash(f'Flashcard added to {topic.name}.', 'success')
    return redirect(url_for('topic_review', topic_id=topic.id))


@app.route('/flashcards/topic/<int:topic_id>')
@login_required
def topic_review(topic_id):
    topic = Topic.query.get_or_404(topic_id)
    if topic.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    cards = Flashcard.query.filter_by(topic_id=topic.id).order_by(Flashcard.created_at).all()
    deck = [{
        'id': c.id,
        'question': c.question,
        'answer': c.answer,
        'review_url': url_for('record_review', card_id=c.id)
    } for c in cards]
    return render_template('topic_review.html', topic=topic, cards=cards, deck=deck)


@app.route('/flashcard/<int:card_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_flashcard(card_id):
    card = Flashcard.query.get_or_404(card_id)
    if card.topic.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    form, subjects = build_flashcard_form()
    if request.method == 'GET':
        form.subject_id.data = card.topic.subject_id
        form.topic_name.data = card.topic.name
        form.question.data = card.question
        form.answer.data = card.answer

    if form.validate_on_submit():
        subject = Subject.query.get(form.subject_id.data)
        if not subject or subject.user_id != current_user.id:
            flash('Access denied.', 'danger')
            return redirect(url_for('dashboard'))

        old_topic = card.topic
        topic = get_or_create_topic(subject, form.topic_name.data)
        card.question = form.question.data.strip()
        card.answer = form.answer.data.strip()
        card.topic_id = topic.id
        db.session.commit()
        if old_topic.id != topic.id and not old_topic.flashcards:
            db.session.delete(old_topic)
            db.session.commit()
        flash('Flashcard updated.', 'success')
        return redirect(url_for('topic_review', topic_id=topic.id))

    flash_form_errors(form)
    return render_template('flashcard_form.html', form=form, card=card, title='Edit Flashcard')


@app.route('/flashcard/<int:card_id>/delete', methods=['POST'])
@login_required
def delete_flashcard(card_id):
    card = Flashcard.query.get_or_404(card_id)
    if card.topic.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    topic = card.topic
    db.session.delete(card)
    db.session.commit()
    if not topic.flashcards:
        db.session.delete(topic)
        db.session.commit()
        flash('Flashcard deleted. That was the last card in that topic, so the topic has gone too.', 'info')
        return redirect(url_for('flashcards_page'))
    flash('Flashcard deleted.', 'info')
    return redirect(url_for('topic_review', topic_id=topic.id))


@app.route('/todo/new', methods=['POST'])
@login_required
def create_todo():
    subjects = Subject.query.filter_by(user_id=current_user.id).order_by(Subject.name).all()
    form = TodoItemForm()
    form.subject_id.choices = [(s.id, s.name) for s in subjects]
    if form.validate_on_submit():
        subject = Subject.query.get(form.subject_id.data)
        if not subject or subject.user_id != current_user.id:
            flash('Access denied.', 'danger')
            return redirect(url_for('dashboard'))
        todo = TodoItem(
            title=form.title.data.strip(),
            scheduled_date=form.scheduled_date.data,
            subject_id=subject.id
        )
        db.session.add(todo)
        db.session.commit()
        flash('To-do added!', 'success')
    else:
        for field_errors in form.errors.values():
            for error in field_errors:
                flash(error, 'danger')
    return redirect(url_for('dashboard'))


@app.route('/todo/<int:todo_id>/toggle', methods=['POST'])
@login_required
def toggle_todo(todo_id):
    todo = TodoItem.query.get_or_404(todo_id)
    if todo.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    todo.status = 'Complete' if todo.status == 'Incomplete' else 'Incomplete'
    db.session.commit()
    return redirect(request.referrer or url_for('dashboard'))


@app.route('/todo/<int:todo_id>/delete', methods=['POST'])
@login_required
def delete_todo(todo_id):
    todo = TodoItem.query.get_or_404(todo_id)
    if todo.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    db.session.delete(todo)
    db.session.commit()
    flash('To-do deleted.', 'info')
    return redirect(url_for('dashboard'))


@app.route('/assessment/<int:assessment_id>/task/new', methods=['POST'])
@login_required
def add_task(assessment_id):
    assessment = Assessment.query.get_or_404(assessment_id)
    if assessment.subject.user_id != current_user.id:
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    form = TaskForm()
    if form.validate_on_submit():
        task = Task(
            title=form.title.data.strip(),
            scheduled_date=form.scheduled_date.data,
            assessment_id=assessment.id
        )
        db.session.add(task)
        db.session.commit()
        flash('Task added!', 'success')
    return redirect(url_for('assessment_detail', assessment_id=assessment.id))

with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=True, port=5000)